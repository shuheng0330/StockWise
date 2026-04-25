from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_DIR = Path(r"C:\Users\User\Documents\UMHackathon\outputs")


def set_cell_shading(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"

    doc.styles["Title"].font.size = Pt(22)
    doc.styles["Subtitle"].font.size = Pt(11)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 3"].font.size = Pt(11)

    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
    else:
        code_style = doc.styles["Code Block"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph(style="Title")
    p.alignment = 1
    p.add_run(title)
    sp = doc.add_paragraph(style="Subtitle")
    sp.alignment = 1
    sp.add_run(subtitle)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_shading(hdr_cells[idx], "CFE2F3")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def build_prd():
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        "StockWise: AI Decision Advisor for Inventory Reordering and Waste Reduction",
        "Revised Product Requirement Document (PRD) for UMHackathon 2026",
    )

    doc.add_heading("1. Product Overview", level=1)
    add_paragraph(
        doc,
        "StockWise is a web-based decision intelligence system for small cafes and kiosks "
        "that need to make daily inventory decisions under waste and cash pressure. The "
        "product converts structured inventory records into ranked actions, simulation-based "
        "trade-off analysis, and Z.AI GLM-powered explanations that help non-technical "
        "operators decide what to restock now, what to buy less of, and what to delay."
    )
    add_paragraph(
        doc,
        "The MVP is grounded in a prototype dataset containing 1,000 records over 100 days, "
        "covering 10 recurring inventory items and 14 structured fields. This allows the "
        "product to demonstrate practical decision support from real inventory behavior "
        "without over-claiming advanced forecasting or revenue prediction."
    )

    doc.add_heading("2. Problem Statement", level=1)
    add_paragraph(
        doc,
        "Small cafes and kiosks often manage inventory through spreadsheets, manual checking, "
        "and intuition. This creates a recurring decision problem: operators know the stock "
        "numbers exist, but they still struggle to determine which items need immediate action, "
        "which items are tying up too much cash, and which purchase quantities will increase "
        "waste. The result is avoidable over-ordering, stockout risk for fast-moving items, "
        "and time lost reviewing raw data."
    )
    add_bullets(
        doc,
        [
            "Operators cannot quickly prioritize which items truly require action.",
            "High-value items can lock up working capital even when demand is uncertain.",
            "Perishable or waste-prone stock creates avoidable losses when ordering is too aggressive.",
            "Manual spreadsheet review is slow and does not explain the trade-offs behind a decision.",
        ],
    )

    doc.add_heading("3. Product Purpose and Business Value", level=1)
    add_paragraph(
        doc,
        "StockWise exists to improve daily reorder decisions. Rather than acting as a passive "
        "dashboard, the system acts as a decision Advisor that combines transparent inventory "
        "logic with AI explanation. The product helps operators convert raw inventory signals "
        "into understandable next-best actions."
    )
    add_bullets(
        doc,
        [
            "Reduce time spent reviewing daily reorder priorities.",
            "Reduce waste-prone over-ordering for perishable or high-risk items.",
            "Reduce cash locked in high-value stock that does not need immediate replenishment.",
            "Increase confidence in reorder actions through explainable, item-specific guidance.",
        ],
    )

    doc.add_heading("4. Target Users", level=1)
    add_paragraph(
        doc,
        "The primary user is a small cafe or kiosk owner-operator who is personally involved "
        "in purchasing and inventory control. This user needs simple, fast, and trustworthy "
        "guidance rather than enterprise software complexity."
    )
    add_table(
        doc,
        ["User", "Needs", "Pain Points"],
        [
            [
                "Small cafe owner / manager",
                "Quick reorder prioritization and cash-aware purchasing decisions",
                "Limited time, spreadsheet dependence, uncertainty about the right order quantity",
            ],
            [
                "Kiosk operator",
                "Simple dashboard and action summary",
                "Difficult to interpret stock levels versus actual usage and lead time",
            ],
            [
                "Kitchen / operations lead",
                "Clear item-level risk flags",
                "Cannot easily explain why an item is urgent or waste-prone",
            ],
        ],
    )

    doc.add_heading("5. Dataset Foundation and Input Scope", level=1)
    add_paragraph(
        doc,
        "The product is intentionally grounded in the dataset that will be used for the demo. "
        "The current prototype dataset includes 100 days of inventory records across 10 recurring "
        "items with no missing values, making it suitable for explainable metric computation, "
        "latest-snapshot analysis, and compact historical context."
    )
    add_table(
        doc,
        ["Input Field", "Role in Decision Support"],
        [
            ["Date", "Supports latest snapshot selection and short historical context"],
            ["Item_ID", "Stable internal item reference"],
            ["Item_Name", "Primary item label shown to users and passed to GLM"],
            ["Category", "Supports grouping and category-level dashboard views"],
            ["Subcategory", "Supports more specific explanation context"],
            ["Unit", "Clarifies quantity interpretation in UI and AI explanation"],
            ["Current_Stock", "Core inventory state metric"],
            ["Reorder_Level", "Reference threshold for urgency assessment"],
            ["Daily_Usage", "Basis for consumption rate and days-of-cover computation"],
            ["Lead_Time", "Supports lead-time demand and stockout-risk logic"],
            ["Price_per_Unit", "Supports cash exposure and waste-cost estimates"],
            ["Supplier_Name", "Adds sourcing context without driving the score directly"],
            ["Seasonal_Factor", "Adjusts demand pressure for time-sensitive items"],
            ["Waste_Percentage", "Supports waste-risk and avoidable-loss analysis"],
        ],
    )

    doc.add_heading("6. Core Product Concept", level=1)
    add_paragraph(
        doc,
        "StockWise is not only a visual analytics screen. It is a decision intelligence workflow "
        "with four linked MVP features that move the user from data to action."
    )
    add_table(
        doc,
        ["MVP Feature", "Description", "Why It Matters"],
        [
            [
                "Decision Dashboard",
                "Shows stock health, low-cover items, waste-cost exposure, and inventory value at risk.",
                "Gives the user a fast operational overview instead of raw rows.",
            ],
            [
                "Priority Action Board",
                "Assigns each item a recommended action: RESTOCK_NOW, BUY_LESS, DELAY_PURCHASE, or MONITOR_CLOSELY.",
                "Turns computed metrics into clear next steps.",
            ],
            [
                "What-If Reorder Simulator",
                "Lets the user test a proposed order quantity and see coverage, cash outlay, and risk change.",
                "Supports trade-off analysis before placing an order.",
            ],
            [
                "GLM Explanation Layer",
                "Uses Z.AI GLM to explain the recommendation in plain language for non-technical users.",
                "Makes the product a true decision-support tool rather than a score table.",
            ],
        ],
    )

    doc.add_heading("7. Feature Details", level=1)
    doc.add_heading("7.1 Decision Dashboard", level=2)
    add_bullets(
        doc,
        [
            "Highlights the latest inventory snapshot for the uploaded CSV.",
            "Shows top urgent items, top waste-cost items, and high inventory value at risk.",
            "Uses category and supplier filters to help users inspect specific segments quickly.",
            "Surfaces compact 7-day context such as usage average and basic trend direction.",
        ],
    )

    doc.add_heading("7.2 Priority Action Board", level=2)
    add_bullets(
        doc,
        [
            "Ranks items using explainable urgency and waste-risk scores.",
            "Presents one recommended action per item with a short reason.",
            "Lets users distinguish between stockout risk and overbuying risk.",
            "Keeps deterministic scores visible even if AI explanation is unavailable.",
        ],
    )

    doc.add_heading("7.3 What-If Reorder Simulator", level=2)
    add_bullets(
        doc,
        [
            "Accepts a proposed reorder quantity for a selected item.",
            "Recomputes simulated coverage days, simulated cash outlay, and simulated risk change.",
            "Helps the user compare aggressive versus conservative ordering before purchase.",
            "Feeds the simulation result into GLM so the explanation reflects the tested scenario.",
        ],
    )

    doc.add_heading("7.4 GLM Explanation Layer", level=2)
    add_bullets(
        doc,
        [
            "Explains the recommended action using only structured metrics supplied by the backend.",
            "Summarizes the trade-off between shortage risk, waste risk, and cash usage.",
            "Provides operator-friendly next steps instead of raw analytical jargon.",
            "Degrades gracefully to rule-based templates if the model output is malformed or unavailable.",
        ],
    )

    doc.add_heading("8. User Stories and Primary Use Cases", level=1)
    add_bullets(
        doc,
        [
            "As a small cafe owner, I want to see which items need attention today so that I do not spend time scanning raw rows.",
            "As a kiosk operator, I want to understand whether a risky item is urgent because of low stock, long lead time, or high waste exposure.",
            "As a purchasing decision-maker, I want to simulate a smaller or larger order quantity so that I can balance stock coverage against waste and cash outlay.",
            "As a non-technical user, I want AI to explain the recommendation in plain language so that I can act confidently without interpreting formulas myself.",
        ],
    )
    add_numbered(
        doc,
        [
            "User uploads the inventory CSV.",
            "StockWise validates the schema and loads the latest snapshot plus short recent context.",
            "The system computes derived metrics and ranks items on the action board.",
            "The user opens a specific item and tests a reorder quantity in the simulator.",
            "The backend recalculates coverage, cash, and risk impact.",
            "Z.AI GLM generates a plain-language explanation based on the structured metrics.",
            "The user chooses an action with a clearer understanding of the trade-off.",
        ],
    )

    doc.add_heading("9. Functional Requirements", level=1)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR1", "The system shall accept an uploaded CSV containing the 14 required inventory fields."],
            ["FR2", "The system shall validate the schema before analysis and reject invalid files with user-friendly feedback."],
            ["FR3", "The system shall compute the latest item-level metrics from the uploaded dataset."],
            ["FR4", "The system shall compute compact recent historical context, including 7-day average usage and trend direction."],
            ["FR5", "The system shall calculate days_of_cover, inventory_value, estimated_waste_cost, lead_time_demand, and stock_gap_to_lead_demand."],
            ["FR6", "The system shall calculate reorder_urgency_score and waste_risk_score using explainable weighted logic."],
            ["FR7", "The system shall classify each item into RESTOCK_NOW, BUY_LESS, DELAY_PURCHASE, or MONITOR_CLOSELY."],
            ["FR8", "The system shall display ranked actions and supporting item metrics in the decision dashboard and action board."],
            ["FR9", "The system shall accept a simulated reorder quantity and recalculate simulated_cash_outlay and simulated_coverage_days."],
            ["FR10", "The system shall pass structured item context to Z.AI GLM and request JSON-only explanation output."],
            ["FR11", "The system shall validate GLM output fields before showing the explanation in the UI."],
            ["FR12", "The system shall retry once with a stricter prompt when GLM returns malformed JSON."],
            ["FR13", "The system shall show deterministic recommendations and a fallback message if GLM output remains invalid or unavailable."],
        ],
    )

    doc.add_heading("10. Non-Functional Requirements", level=1)
    add_bullets(
        doc,
        [
            "Usability: the UI must remain readable and easy to interpret for non-technical operators.",
            "Explainability: each recommendation must be traceable to visible metrics or AI explanation derived from those metrics.",
            "Performance: the app should analyze a prototype-sized CSV and produce the first dashboard response in seconds, not minutes.",
            "Reliability: the system must handle missing columns, invalid files, and AI API issues without crashing.",
            "Maintainability: the backend should keep validation, metric computation, recommendation logic, simulation, and AI integration as separate modules.",
        ],
    )

    doc.add_heading("11. Derived Metrics and Decision Logic", level=1)
    add_paragraph(
        doc,
        "The product uses transparent derived metrics instead of opaque forecasting. These metrics convert raw inventory fields into a decision-ready view."
    )
    add_table(
        doc,
        ["Derived Metric", "Definition", "Purpose"],
        [
            ["days_of_cover", "Current_Stock / Daily_Usage", "Indicates how many days current stock can support"],
            ["inventory_value", "Current_Stock * Price_per_Unit", "Shows cash tied up in current stock"],
            ["estimated_waste_cost", "inventory_value * (Waste_Percentage / 100)", "Estimates avoidable loss exposure"],
            ["lead_time_demand", "Daily_Usage * Lead_Time * Seasonal_Factor", "Estimates stock needed before replenishment arrives"],
            ["stock_gap_to_lead_demand", "Current_Stock - lead_time_demand", "Shows whether stock is likely to cover lead-time demand"],
            ["simulated_coverage_days", "(Current_Stock + simulated_order_qty) / Daily_Usage", "Shows coverage after a tested reorder quantity"],
            ["simulated_cash_outlay", "simulated_order_qty * Price_per_Unit", "Shows cash needed for a tested reorder"],
        ],
    )
    add_paragraph(
        doc,
        "Reorder urgency is defined by multiple factors, not reorder threshold alone. The score emphasizes low coverage relative to lead-time demand, reorder threshold proximity, usage pressure, lead time, and seasonal pressure. Waste risk emphasizes waste percentage, inventory value, and over-coverage risk. This keeps the MVP explainable and aligned to the available dataset."
    )

    doc.add_heading("12. Z.AI GLM Integration Strategy", level=1)
    doc.add_heading("12.1 Why Z.AI GLM Is Essential", level=2)
    add_paragraph(
        doc,
        "Z.AI GLM is the intelligence layer that transforms structured metrics into operator-friendly reasoning. "
        "Without GLM, StockWise can still calculate scores and labels, but it degrades into a monitoring dashboard. "
        "The product would lose the human-readable reasoning, trade-off explanation, and next-step guidance that convert raw analytics into meaningful decision support."
    )

    doc.add_heading("12.2 Prompting Strategy", level=2)
    add_paragraph(
        doc,
        "The system uses a bounded structured prompting strategy. Deterministic logic computes the metrics and recommendation first. "
        "GLM is then asked to explain, not invent, the decision. This reduces hallucination risk and keeps the explanation anchored to visible facts."
    )
    add_paragraph(doc, "System prompt:")
    add_code_block(
        doc,
        'You are StockWise, an inventory decision Advisor for small cafe operators. '
        'Use only the provided inventory metrics and structured findings. Explain the recommended action clearly and briefly. '
        'Do not invent sales, profit, or supplier facts that are not in the input. Return valid JSON only.',
    )
    add_paragraph(doc, "User prompt pattern:")
    add_code_block(
        doc,
        "{\n"
        '  "task": "Explain the recommended inventory action for a small cafe operator.",\n'
        '  "item_context": {\n'
        '    "item_name": "Paneer",\n'
        '    "category": "Veg",\n'
        '    "subcategory": "Dairy",\n'
        '    "unit": "kg",\n'
        '    "supplier_name": "Supplier A"\n'
        "  },\n"
        '  "current_metrics": {\n'
        '    "current_stock": 12.54,\n'
        '    "reorder_level": 5.21,\n'
        '    "daily_usage": 2.24,\n'
        '    "lead_time": 4,\n'
        '    "seasonal_factor": 1.30,\n'
        '    "price_per_unit": 450,\n'
        '    "waste_percentage": 4.74,\n'
        '    "days_of_cover": 5.60,\n'
        '    "inventory_value": 5643.00,\n'
        '    "estimated_waste_cost": 267.48,\n'
        '    "lead_time_demand": 11.65,\n'
        '    "stock_gap_to_lead_demand": 0.89,\n'
        '    "reorder_urgency_score": 72,\n'
        '    "waste_risk_score": 81,\n'
        '    "recommended_action": "MONITOR_CLOSELY"\n'
        "  },\n"
        '  "recent_context": {\n'
        '    "avg_usage_7d": 2.37,\n'
        '    "trend_direction": "up"\n'
        "  },\n"
        '  "simulation_context": {\n'
        '    "simulated_order_qty": 3.0,\n'
        '    "simulated_cash_outlay": 1350.0,\n'
        '    "simulated_coverage_days": 6.94,\n'
        '    "simulated_risk_change": "higher waste exposure with limited urgency reduction"\n'
        "  }\n"
        "}",
    )

    doc.add_heading("12.3 Structured Context Passed to GLM", level=2)
    add_bullets(
        doc,
        [
            "Item identity: item_name, category, subcategory, unit, supplier_name",
            "Current-state metrics: current_stock, reorder_level, daily_usage, lead_time, seasonal_factor, price_per_unit, waste_percentage",
            "Derived metrics: days_of_cover, inventory_value, estimated_waste_cost, lead_time_demand, stock_gap_to_lead_demand, reorder_urgency_score, waste_risk_score, recommended_action",
            "Optional recent context: avg_usage_7d and trend_direction",
            "Optional simulation context: simulated_order_qty, simulated_cash_outlay, simulated_coverage_days, simulated_risk_change",
        ],
    )
    add_paragraph(
        doc,
        "Oversized input is controlled by sending only the latest item snapshot plus compact historical summaries, not the raw 1,000-row dataset. This keeps the GLM context bounded and aligned with the competition requirement to explain how input size is managed."
    )

    doc.add_heading("12.4 Expected Response Schema", level=2)
    add_table(
        doc,
        ["Field", "Requirement"],
        [
            ["item_name", "Must match the requested item"],
            ["recommended_action", "One of RESTOCK_NOW, BUY_LESS, DELAY_PURCHASE, MONITOR_CLOSELY"],
            ["priority_level", "One of HIGH, MEDIUM, LOW"],
            ["short_reason", "One-sentence explanation summary"],
            ["decision_explanation", "Short operator-facing rationale"],
            ["tradeoff_summary", "Explains shortage vs waste vs cash trade-off"],
            ["suggested_next_step", "Practical action to take next"],
            ["confidence_note", "Clarifies confidence or limitation"],
            ["warning_flag", "Optional note if risk remains high or context is incomplete"],
        ],
    )

    doc.add_heading("12.5 Response Parsing and Validation", level=2)
    add_numbered(
        doc,
        [
            "The backend sends the structured prompt to Z.AI GLM and expects JSON output only.",
            "The response parser checks that all required keys exist.",
            "The parser validates enum values for recommended_action and priority_level.",
            "The parser checks explanation length and ensures the item_name matches the request context.",
            "The parser rejects unsupported claims such as revenue prediction or profit optimization that are not justified by the dataset.",
            "If valid, the explanation is shown in the UI together with the deterministic scores.",
            "If partially valid, safe fields are kept and bad or missing fields are replaced with rule-based templates.",
            "If invalid, the system triggers the fallback flow and keeps the dashboard recommendations visible.",
        ],
    )

    doc.add_heading("12.6 Graceful Fallback Behavior", level=2)
    add_bullets(
        doc,
        [
            "Level 1: Retry once with a stricter JSON-only prompt if the first response is malformed.",
            "Level 2: If the retry still fails, use a deterministic explanation template such as: Restock Paneer now because days of cover is low relative to lead-time demand and the item has high inventory value.",
            "Level 3: Keep dashboard scores, rankings, and recommended actions visible even when AI explanation is unavailable.",
            "User-safe message: AI explanation is temporarily unavailable. StockWise is showing rule-based recommendations instead.",
        ],
    )

    doc.add_heading("13. Success Metrics", level=1)
    add_bullets(
        doc,
        [
            "Users can identify the top priority reorder items in under two minutes from the dashboard view.",
            "Users can compare at least two reorder quantity options in the simulator before deciding.",
            "The system surfaces high-value or high-waste exposure items clearly enough to justify action during a demo walkthrough.",
            "The app demonstrates lower manual review effort than scanning the raw CSV directly.",
            "The product clearly shows GLM as essential for explanation and decision guidance, not as decorative text generation.",
        ],
    )

    doc.add_heading("14. Out of Scope", level=1)
    add_bullets(
        doc,
        [
            "Supplier ordering automation and purchase-order execution",
            "Revenue or profit forecasting from sales transactions",
            "Enterprise ERP integration",
            "Multi-branch stock synchronization",
            "IoT real-time stock sensing",
            "Full mobile application support in the MVP round",
        ],
    )

    doc.add_heading("15. Assumptions, Constraints, and Risks", level=1)
    add_bullets(
        doc,
        [
            "The prototype dataset is representative enough to demonstrate operational decision support.",
            "The primary workflow uses the latest snapshot plus compact recent history, not advanced forecasting.",
            "The hackathon scope favors explainable rules and simulation over opaque machine learning.",
            "The product must use Z.AI GLM to remain eligible for judging.",
            "If recommendations feel generic, the demo value drops; this is mitigated by grounding explanations in visible metrics and item-specific examples such as Mutton, Paneer, and Eggs.",
        ],
    )

    doc.add_heading("16. Conclusion", level=1)
    add_paragraph(
        doc,
        "StockWise aligns with the competition’s decision-intelligence domain by showing how AI helps operators make smarter, data-informed inventory decisions with practical economic value. "
        "The revised MVP is intentionally achievable as a working web app while still demonstrating originality through simulation, trade-off-aware recommendations, and Z.AI GLM-powered explanation."
    )
    return doc


def build_sad():
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        "StockWise: AI Decision Advisor for Inventory Reordering and Waste Reduction",
        "Revised System Analysis Document (SAD) for UMHackathon 2026",
    )

    doc.add_heading("1. Purpose and System Overview", level=1)
    add_paragraph(
        doc,
        "StockWise is a web-based inventory decision support system for small cafes and kiosks. "
        "Its purpose is to transform uploaded inventory records into explainable actions by combining "
        "deterministic metric computation, action classification, reorder simulation, and Z.AI GLM explanation."
    )
    add_paragraph(
        doc,
        "The system is designed as a hackathon-grade working MVP. It uses a structured CSV input, computes decision-ready metrics for the latest snapshot and recent context, then produces a ranked action board and item-level explanation. The architecture is intentionally modular so the product appears technically credible and extensible."
    )

    doc.add_heading("2. Stakeholders and System Goals", level=1)
    add_table(
        doc,
        ["Stakeholder", "Interest"],
        [
            ["Small cafe / kiosk operator", "Needs fast and explainable reorder decisions"],
            ["Hackathon judges", "Need evidence of decision intelligence, GLM centrality, and technical feasibility"],
            ["Development team", "Needs a scoped architecture that can be built and demoed reliably"],
        ],
    )
    add_bullets(
        doc,
        [
            "Compute transparent inventory metrics from the uploaded dataset.",
            "Classify each item into an action that supports a real purchasing decision.",
            "Allow users to simulate reorder quantities before acting.",
            "Use Z.AI GLM as a service layer that explains recommendations and trade-offs.",
        ],
    )

    doc.add_heading("3. Scope and Architectural Principles", level=1)
    add_bullets(
        doc,
        [
            "Scope is limited to a working web app for inventory decision support, not full stock management automation.",
            "Business logic remains deterministic and explainable.",
            "GLM is used to convert structured findings into operator-friendly reasoning.",
            "Historical data is used only for compact trend support, not heavy forecasting claims.",
            "The architecture separates validation, metrics, recommendation logic, simulation, and AI integration to improve clarity and maintainability.",
        ],
    )

    doc.add_heading("4. Data Source and Schema", level=1)
    add_paragraph(
        doc,
        "The system uses the demo CSV dataset as the authoritative input source. The current dataset contains 1,000 rows over 100 dates for 10 recurring items, with 14 structured fields and no missing values detected in the prototype analysis."
    )
    add_table(
        doc,
        ["Field", "Type", "System Use"],
        [
            ["Date", "date/string", "Latest snapshot and recent trend context"],
            ["Item_ID", "integer", "Stable key"],
            ["Item_Name", "string", "Display label and AI reference"],
            ["Category", "string", "Dashboard grouping"],
            ["Subcategory", "string", "Detailed grouping"],
            ["Unit", "string", "Quantity interpretation"],
            ["Current_Stock", "float", "Current availability"],
            ["Reorder_Level", "float", "Reference threshold"],
            ["Daily_Usage", "float", "Consumption rate"],
            ["Lead_Time", "integer", "Supply delay factor"],
            ["Price_per_Unit", "numeric", "Cash and waste-cost exposure"],
            ["Supplier_Name", "string", "Source context"],
            ["Seasonal_Factor", "float", "Demand adjustment"],
            ["Waste_Percentage", "float", "Waste-loss exposure"],
        ],
    )

    doc.add_heading("5. High-Level Architecture", level=1)
    add_paragraph(doc, "Diagram-ready flow:")
    add_code_block(
        doc,
        "User Interface -> Backend API -> CSV Validation -> Metric Engine -> Recommendation Engine -> "
        "Simulation Engine -> Z.AI GLM Adapter -> Response Parser/Fallback -> Dashboard + Action Board + Explanation Panel"
    )
    add_table(
        doc,
        ["Component", "Responsibility", "Output"],
        [
            ["Frontend UI", "Upload CSV, display KPIs, rankings, simulator, and explanation", "User-facing dashboard"],
            ["CSV Validation Layer", "Check schema, data types, required columns, empty files", "Validated structured dataset or user-safe error"],
            ["Metric Engine", "Compute derived metrics and recent context", "Item-level analytical features"],
            ["Recommendation Engine", "Score urgency and waste risk, assign action label", "Ranked actions and KPI summaries"],
            ["Simulation Engine", "Recompute metrics after user-entered reorder quantity", "Scenario-based impact values"],
            ["GLM Adapter", "Build prompt payload and call Z.AI GLM", "Raw model response"],
            ["Parser/Fallback Layer", "Validate JSON response and recover gracefully from failures", "Safe explanation payload"],
        ],
    )

    doc.add_heading("6. Core Metrics and Decision Logic", level=1)
    add_paragraph(
        doc,
        "The analytical core is deterministic. Scores are computed before any AI call so the recommendation remains auditable."
    )
    add_table(
        doc,
        ["Metric", "Formula", "Interpretation"],
        [
            ["days_of_cover", "Current_Stock / Daily_Usage", "Estimated days current stock can support"],
            ["inventory_value", "Current_Stock * Price_per_Unit", "Cash tied up in stock"],
            ["estimated_waste_cost", "inventory_value * (Waste_Percentage / 100)", "Approximate avoidable waste loss"],
            ["lead_time_demand", "Daily_Usage * Lead_Time * Seasonal_Factor", "Projected quantity required before next replenishment"],
            ["stock_gap_to_lead_demand", "Current_Stock - lead_time_demand", "Negative values indicate likely shortfall pressure"],
            ["simulated_coverage_days", "(Current_Stock + simulated_order_qty) / Daily_Usage", "Coverage if the tested order is placed"],
            ["simulated_cash_outlay", "simulated_order_qty * Price_per_Unit", "Cash required for the tested order"],
        ],
    )
    add_paragraph(
        doc,
        "Reorder urgency is intentionally based on multiple signals because the dataset shows that only a small percentage of rows sit below the reorder threshold alone. The engine therefore emphasizes coverage relative to lead-time demand, reorder threshold proximity, usage intensity, lead time, and seasonal pressure."
    )
    add_bullets(
        doc,
        [
            "Reorder urgency score: weighted combination of low coverage, lead-time pressure, reorder threshold pressure, usage intensity, and seasonal pressure.",
            "Waste risk score: weighted combination of waste percentage, inventory value, and over-coverage or delayed-turnover risk.",
            "Action classification remains rule-based and explainable so it can be defended to judges.",
        ],
    )

    doc.add_heading("7. Recommendation Classification", level=1)
    add_table(
        doc,
        ["Action", "Typical Trigger"],
        [
            [
                "RESTOCK_NOW",
                "High urgency because coverage is weak relative to lead-time demand or stock gap is negative",
            ],
            [
                "BUY_LESS",
                "High waste or cash exposure with enough current coverage, so replenishment should be conservative",
            ],
            [
                "DELAY_PURCHASE",
                "Current coverage is strong and risk is low, so immediate purchase is unnecessary",
            ],
            [
                "MONITOR_CLOSELY",
                "Item is not critical yet but signals are rising and should be rechecked soon",
            ],
        ],
    )
    add_paragraph(
        doc,
        "Example narrative supported by the current dataset: high-value items such as Mutton or Paneer may not always require immediate replenishment, but they can still represent significant waste-cost exposure. Fast-moving items such as Eggs may have lower unit cost but still need attention when coverage becomes tight."
    )

    doc.add_heading("8. AI Service Layer Integration", level=1)
    doc.add_heading("8.1 Role of Z.AI GLM", level=2)
    add_paragraph(
        doc,
        "Z.AI GLM is integrated as a service layer, not as a vague AI box. The model does not calculate raw scores. Instead, it receives a structured decision context from the backend and produces a JSON explanation that clarifies why the recommended action makes sense for a small cafe operator."
    )
    add_paragraph(
        doc,
        "If GLM is removed, StockWise can still display scores, but it loses the explanation, trade-off summary, and natural-language action guidance that make the system effective as a decision-support product. This keeps GLM central to user value while preserving deterministic auditability."
    )

    doc.add_heading("8.2 Prompt Sent to GLM", level=2)
    add_paragraph(doc, "System prompt:")
    add_code_block(
        doc,
        'You are StockWise, an inventory decision Advisor for small cafe operators. '
        'Use only the provided inventory metrics and structured findings. Explain the recommended action clearly and briefly. '
        'Do not invent sales, profit, or supplier facts that are not in the input. Return valid JSON only.'
    )
    add_paragraph(doc, "User prompt content:")
    add_paragraph(
        doc,
        "The backend sends a structured object containing item identity, current metrics, derived metrics, recent context, and optional simulation context. The prompt asks GLM to explain the recommended action, mention the shortage-versus-waste-versus-cash trade-off, and keep the answer concise for an operator-facing panel."
    )

    doc.add_heading("8.3 Structured Context Included", level=2)
    add_table(
        doc,
        ["Context Group", "Fields"],
        [
            ["Item identity", "item_name, category, subcategory, unit, supplier_name"],
            ["Current-state metrics", "current_stock, reorder_level, daily_usage, lead_time, seasonal_factor, price_per_unit, waste_percentage"],
            ["Derived metrics", "days_of_cover, inventory_value, estimated_waste_cost, lead_time_demand, stock_gap_to_lead_demand, reorder_urgency_score, waste_risk_score, recommended_action"],
            ["Recent context", "avg_usage_7d, trend_direction"],
            ["Simulation context", "simulated_order_qty, simulated_cash_outlay, simulated_coverage_days, simulated_risk_change"],
        ],
    )
    add_paragraph(
        doc,
        "Oversized input is prevented by design. The system does not send the full 1,000-row CSV to GLM. It sends only the latest item snapshot and compact historical summaries, which keeps the context window bounded and predictable."
    )

    doc.add_heading("8.4 Expected Response Fields", level=2)
    add_code_block(
        doc,
        "{\n"
        '  "item_name": "Paneer",\n'
        '  "recommended_action": "MONITOR_CLOSELY",\n'
        '  "priority_level": "MEDIUM",\n'
        '  "short_reason": "Paneer is close to lead-time demand and carries high waste-cost exposure.",\n'
        '  "decision_explanation": "Current stock is still slightly above projected lead-time demand, so immediate restocking is not required. Because Paneer is expensive and waste-prone, ordering more right now would increase avoidable cash and waste exposure.",\n'
        '  "tradeoff_summary": "Delaying a large order lowers waste and cash risk, but the item should be reviewed again soon because demand is trending up.",\n'
        '  "suggested_next_step": "Monitor for one cycle or place only a small top-up order.",\n'
        '  "confidence_note": "Confidence is moderate because recent usage trend is increasing.",\n'
        '  "warning_flag": "High-value item with elevated waste percentage."\n'
        "}"
    )

    doc.add_heading("8.5 Response Parsing and Validation", level=2)
    add_numbered(
        doc,
        [
            "Backend sends the structured prompt to Z.AI GLM.",
            "Backend expects JSON-only output.",
            "Parser checks that all required keys exist.",
            "Parser validates enum values for recommended_action and priority_level.",
            "Parser checks that item_name matches the current request context.",
            "Parser enforces length limits so explanations stay concise in the UI.",
            "Parser rejects unsupported claims such as revenue prediction or profit optimization that are not supported by the dataset.",
            "If the response passes validation, the explanation payload is returned to the frontend.",
            "If the response is partially valid, safe fields are retained and missing or unsafe text is replaced with deterministic templates.",
            "If the response is invalid, fallback logic is triggered and the deterministic recommendation remains visible.",
        ],
    )

    doc.add_heading("8.6 Graceful Fallback Behavior", level=2)
    add_bullets(
        doc,
        [
            "Level 1 retry: send a stricter JSON-only prompt if the first response is malformed.",
            "Level 2 deterministic template: if still invalid, generate a safe explanation such as Restock Paneer now because days of cover is low relative to lead-time demand and the item has high inventory value.",
            "Level 3 safe UI state: keep scores, rankings, and recommended actions visible with a clear fallback message.",
            "Fallback message example: AI explanation is temporarily unavailable. StockWise is showing rule-based recommendations instead.",
        ],
    )

    doc.add_heading("9. End-to-End User Journey", level=1)
    add_numbered(
        doc,
        [
            "The user uploads the inventory CSV through the web interface.",
            "The backend validates required columns and data quality.",
            "The latest snapshot is selected and recent context is derived from the last seven days per item.",
            "The metric engine computes core metrics and scores.",
            "The recommendation engine assigns an action label and builds ranked outputs.",
            "The frontend displays the dashboard and priority action board.",
            "The user selects an item and enters a simulated reorder quantity.",
            "The simulation engine recalculates coverage, cash outlay, and risk change.",
            "The backend sends the bounded structured context to GLM.",
            "The parser validates the model response and either returns the explanation or triggers fallback.",
            "The frontend shows the final action rationale and simulation impact to the user.",
        ],
    )

    doc.add_heading("10. Backend Contract", level=1)
    add_table(
        doc,
        ["Operation", "Purpose", "Request", "Response"],
        [
            [
                "Upload and analyze inventory",
                "Validate the CSV and compute current item metrics",
                "CSV file containing the 14 required fields",
                "KPI summary, ranked items, derived metrics, and recommended_action per item",
            ],
            [
                "Request item simulation",
                "Test a proposed reorder quantity for one item",
                "item_id or item_name plus simulated_order_qty",
                "Updated simulated_coverage_days, simulated_cash_outlay, simulated_risk_change, and current scores",
            ],
            [
                "Request GLM explanation",
                "Generate operator-facing explanation for one item decision",
                "Structured item context, current metrics, derived metrics, and optional simulation context",
                "Validated explanation payload or fallback-safe explanation",
            ],
        ],
    )

    doc.add_heading("11. Error Handling and Resilience", level=1)
    add_bullets(
        doc,
        [
            "Invalid CSV schema: reject file and show missing required columns.",
            "Empty file or unreadable rows: stop analysis and show user-safe error guidance.",
            "Division edge cases such as near-zero usage: guard formulas to avoid invalid coverage values.",
            "GLM timeout or malformed response: activate retry and fallback without hiding deterministic outputs.",
            "Unexpected backend exception: preserve safe UI state and avoid exposing sensitive implementation details.",
        ],
    )

    doc.add_heading("12. Security, Privacy, and Operational Constraints", level=1)
    add_bullets(
        doc,
        [
            "API keys must be stored in environment variables rather than hardcoded in the client or repository.",
            "Uploaded files are used for analysis only and should not expose secrets or infrastructure details.",
            "The MVP assumes prototype-scale usage and does not claim enterprise security controls.",
            "The judging requirement that Z.AI is mandatory is treated as a hard dependency for the demo environment.",
        ],
    )

    doc.add_heading("13. Assumptions, Dependencies, and Future Enhancements", level=1)
    add_bullets(
        doc,
        [
            "The prototype dataset remains the authoritative source for the MVP and is sufficient for demonstrating decision support.",
            "Recent historical context remains compact and explainable rather than forecast-heavy.",
            "Frontend, backend, and GLM API connectivity are available during the demo.",
            "Future enhancements may include supplier-aware reorder workflows, branch comparisons, alerting, and richer forecasting once sales or demand data exists.",
        ],
    )

    doc.add_heading("14. Conclusion", level=1)
    add_paragraph(
        doc,
        "The revised StockWise architecture demonstrates a credible decision-intelligence product for the competition. "
        "It shows how structured data, transparent scoring, bounded simulation, and Z.AI GLM explanation can work together as a practical system for economic decision support. "
        "The design stays ambitious enough to score on originality and user value while remaining scoped to a buildable hackathon MVP."
    )
    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    prd_path = OUTPUT_DIR / "StockWise_PRD_Revised.docx"
    sad_path = OUTPUT_DIR / "StockWise_SAD_Revised.docx"

    build_prd().save(prd_path)
    build_sad().save(sad_path)

    print(prd_path)
    print(sad_path)


if __name__ == "__main__":
    main()
